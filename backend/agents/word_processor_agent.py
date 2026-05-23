"""
agents/word_processor_agent.py
───────────────────────────────
Extracts readable text from a base64-encoded Word document (.docx) and
returns it as a plain string for context injection by the AgentOrchestrator.

Supports .docx only (Office Open XML).  The legacy .doc binary format is not
supported — browsers and modern Office versions all produce .docx.

Text is extracted in document order: paragraphs first, then table cells.
For very long documents the result is truncated to _MAX_CHARS.

Dependency: python-docx  (pip install python-docx)
"""

from __future__ import annotations

import base64
import io
import logging

logger = logging.getLogger(__name__)

# Hard cap on extracted text injected into the specialist agent's context.
# ~16 000 chars ≈ ~4 000 tokens — enough for a dense 10-12 page document
# or a detailed multi-role resume.
_MAX_CHARS = 16_000


class WordProcessorAgent:
    """
    Extracts text from a base64-encoded .docx file.

    Usage
    ─────
        agent = WordProcessorAgent()
        text = await agent.process(docx_base64)
        # text is plain extracted content, ready to inject into the specialist prompt
    """

    async def process(self, docx_base64: str) -> str:
        """
        Decode the base64 .docx and extract all readable text.

        Returns the extracted text on success, or an error note on failure.
        Truncates at _MAX_CHARS with a notice when the document is very long.
        """
        try:
            # Import here so a missing python-docx doesn't break other agents
            # at import time — only raises when a Word file is actually submitted.
            import docx  # noqa: PLC0415

            doc_bytes = base64.b64decode(docx_base64)
            document = docx.Document(io.BytesIO(doc_bytes))

            sections: list[str] = []

            # ── Paragraphs (body text, headings, list items) ─────────────────
            para_texts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
            if para_texts:
                sections.append("\n\n".join(para_texts))

            # ── Tables ───────────────────────────────────────────────────────
            for table in document.tables:
                rows: list[str] = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                if rows:
                    sections.append("\n".join(rows))

            full_text = "\n\n".join(sections)

            if not full_text.strip():
                return "[Word document contained no extractable text.]"

            truncated = False
            if len(full_text) > _MAX_CHARS:
                full_text = full_text[:_MAX_CHARS]
                truncated = True

            para_count = len(para_texts)
            table_count = len(document.tables)
            logger.info(
                "WordProcessorAgent: extracted %d chars (%d paragraphs, %d tables)%s",
                len(full_text),
                para_count,
                table_count,
                " (truncated)" if truncated else "",
            )

            result = (
                f"**Extracted Word document content "
                f"({para_count} paragraph(s), {table_count} table(s)):**\n\n"
                f"{full_text}"
            )
            if truncated:
                result += f"\n\n[Document truncated at {_MAX_CHARS} characters]"
            return result

        except ImportError:
            logger.error("WordProcessorAgent: python-docx is not installed")
            return "[Word document processing is unavailable: python-docx package not installed.]"
        except Exception as exc:
            logger.error(
                "WordProcessorAgent.process failed: %s", exc, exc_info=True
            )
            return f"[Word document could not be processed: {exc}]"
